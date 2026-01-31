# services/word_services.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.conf import settings
import os


def add_center_text(doc, text, size=14, bold=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(size)
    run.bold = bold


def build_docx_academic(data, filename):
    os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
    doc = Document()

    # ==== TITUL BET ====
    add_center_text(doc, data["ministry"], 12, True)
    add_center_text(doc, data["university"], 12, True)
    add_center_text(doc, data["department"], 12, True)

    doc.add_paragraph("\n")

    add_center_text(doc, data["work_type"], 14, True)
    add_center_text(doc, f"MAVZU: {data['title']}", 14, True)

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.add_run(f"TAYYORLADI: {data['author']}\n").bold = True
    p.add_run(data["course"])

    doc.add_paragraph("\n")
    doc.add_paragraph(f"TEKSHIRDI: {data['checker']}")

    doc.add_paragraph("\n\n")
    add_center_text(doc, data["city_year"], 12)

    doc.add_page_break()

    # ==== REJA ====
    doc.add_heading("REJA", level=1)
    for i, item in enumerate(data["plan"], start=1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    doc.add_page_break()

    # ==== ASOSIY BO‘LIMLAR ====
    for section in data["sections"]:
        doc.add_heading(section["heading"], level=2)
        p = doc.add_paragraph(section["content"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # ==== KOD MISOLLAR ====
    if "code_examples" in data:
        doc.add_heading("Amaliy misollar", level=2)
        for example in data["code_examples"]:
            doc.add_paragraph(example["description"], style="List Bullet")
            code_para = doc.add_paragraph(example["code"])
            code_para.paragraph_format.left_indent = Pt(24)

    # ==== XULOSA ====
    doc.add_heading("Xulosa", level=2)
    p = doc.add_paragraph(data["conclusion"])
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    path = os.path.join(settings.MEDIA_ROOT, f"{filename}.docx")
    doc.save(path)

    return settings.MEDIA_URL + f"{filename}.docx"